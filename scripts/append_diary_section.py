#!/usr/bin/env python3
"""
append_diary_section.py - Inserisce nel diario .docx le sezioni di un draft Markdown.

Il diario tecnico e' un .docx con formattazione ricca, e per questo la sua
manutenzione e' sempre stata manuale: si apriva Word, si incollava a mano il
draft prodotto in scratchpad, si riformattava. Il costo di quel passo e' la
ragione per cui il documento ha accumulato due mesi di ritardo.

Questo script lo automatizza per il caso che copre la quasi totalita' degli
aggiornamenti reali, cioe' l'aggiunta in coda di nuove sezioni di ciclo. Non
riscrive nulla di esistente: inserisce i paragrafi nuovi immediatamente prima
di un paragrafo di ancoraggio, per default l'intestazione "Lezioni apprese",
usando gli stili gia' presenti nel documento (Heading 2, Heading 3, Normal) e
convertendo le note del draft in note a pie' di pagina vere di Word.

Formato del draft atteso (lo stesso che l'agente produce in scratchpad):

    ## C.9 Titolo della sezione      -> Heading 2
    ### Sottosezione                 -> Heading 3
    testo del paragrafo              -> Normal
    *termine denso*                  -> corsivo
    `keyword`                        -> monospazio
    **testo**                        -> grassetto
    qualcosa[^1]                     -> riferimento a nota a pie' di pagina
    [^1]: testo della nota           -> corpo della nota (in fondo al draft)

Tutto cio' che precede la prima riga `## ` viene ignorato, cosi' il draft puo'
aprirsi con un titolo e delle istruzioni per il lettore umano senza che
finiscano nel documento.

Uso:
  # cosa verrebbe inserito, senza scrivere nulla
  python scripts/append_diary_section.py --draft _notes/draft-diario-C9-C12.md

  # inserimento reale (crea sempre un .bak prima di toccare il file)
  python scripts/append_diary_section.py --draft _notes/draft-diario-C9-C12.md --apply

Dopo l'inserimento va eseguito `.\\scripts\\finalize_diary.ps1`, che rigenera il
.md e mostra il diff: quella resta la review vera del risultato.
"""

import argparse
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
except ImportError:
    print("ERRORE: python-docx non installato. Esegui: pip install python-docx",
          file=sys.stderr)
    sys.exit(1)


DIARY_NAME = "diario-tecnico-progetto (lettore-doc + skills-repo).docx"
DEFAULT_ANCHOR = "Lezioni apprese"

# Stile di carattere del rimando alla nota, come gia' presente nel documento.
# Il nome e' quello italiano generato da Word: se un giorno il diario venisse
# ricreato con un Word in lingua diversa questo valore andrebbe riletto dal
# footnotes.xml invece che assunto.
FOOTNOTE_REF_STYLE = "Rimandonotaapidipagina"

# Font monospazio per le keyword di codice inline.
MONO_FONT = "Consolas"


# ---------------------------------------------------------------------------
# Parsing del draft
# ---------------------------------------------------------------------------

FOOTNOTE_DEF_RE = re.compile(r"^\[\^([0-9]+)\]:\s*(.+)$")
HEADING_RE      = re.compile(r"^(#{1,4})\s+(.*)$")

# Tokenizzatore inline: cattura `code`, **bold**, *italic*, [^n].
INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|\[\^[0-9]+\])")


def parse_draft(text: str) -> tuple[list[dict], dict[str, str]]:
    """
    Restituisce (blocchi, note).

    Ogni blocco e' {"kind": "heading"|"para", "level": int, "text": str}.
    Le note sono {numero_locale: testo}.
    """
    footnotes: dict[str, str] = {}
    body_lines: list[str] = []

    # Prima passata: estrae le definizioni di nota, ovunque siano.
    for line in text.splitlines():
        m = FOOTNOTE_DEF_RE.match(line.strip())
        if m:
            footnotes[m.group(1)] = m.group(2).strip()
            continue
        body_lines.append(line)

    # Le definizioni possono continuare sulla riga successiva: si ricuciono
    # unendo al testo della nota le righe che la seguono fino alla riga vuota.
    lines = text.splitlines()
    for i, line in enumerate(lines):
        m = FOOTNOTE_DEF_RE.match(line.strip())
        if not m:
            continue
        extra: list[str] = []
        for follow in lines[i + 1:]:
            if not follow.strip() or FOOTNOTE_DEF_RE.match(follow.strip()):
                break
            extra.append(follow.strip())
            if follow in body_lines:
                body_lines.remove(follow)
        if extra:
            footnotes[m.group(1)] = footnotes[m.group(1)] + " " + " ".join(extra)

    # Seconda passata: si parte dalla prima intestazione di sezione, cosi' il
    # preambolo del draft (titolo e istruzioni per l'umano) resta fuori.
    blocks: list[dict] = []
    started = False
    buffer: list[str] = []

    def flush():
        if buffer:
            joined = " ".join(s.strip() for s in buffer).strip()
            if joined:
                blocks.append({"kind": "para", "level": 0, "text": joined})
            buffer.clear()

    for line in body_lines:
        stripped = line.strip()
        heading = HEADING_RE.match(stripped)

        if heading and len(heading.group(1)) >= 2:
            started = True
        if not started:
            continue

        if heading:
            flush()
            blocks.append({
                "kind": "heading",
                "level": len(heading.group(1)),
                "text": heading.group(2).strip(),
            })
        elif not stripped or stripped == "---":
            flush()
        else:
            buffer.append(stripped)

    flush()
    return blocks, footnotes


# ---------------------------------------------------------------------------
# Scrittura nel documento
# ---------------------------------------------------------------------------

def footnotes_part(document):
    """La parte word/footnotes.xml, che python-docx tratta come blob opaco."""
    for rel in document.part.rels.values():
        if "footnotes" in rel.reltype:
            return rel.target_part
    return None


def next_footnote_id(xml: str) -> int:
    """Primo id libero in footnotes.xml (le note reali partono da 1)."""
    ids = [int(v) for v in re.findall(r'<w:footnote[^>]*w:id="(-?\d+)"', xml)]
    return max(ids) + 1 if ids else 1


def xml_escape(text: str) -> str:
    return (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def build_footnote_xml(fid: int, text: str) -> str:
    """
    Una nota nuova, con lo stesso impianto di quelle gia' nel documento.

    Il testo passa per lo stesso tokenizzatore inline del corpo, cosi' il
    corsivo con cui si marca l'acronimo e il monospazio delle keyword vengono
    resi come formattazione vera invece di finire nella nota come asterischi
    e backtick letterali.
    """
    base = '<w:color w:val="595959"/><w:sz w:val="18"/><w:szCs w:val="18"/>'
    runs = [
        f'<w:r><w:rPr><w:rStyle w:val="{FOOTNOTE_REF_STYLE}"/></w:rPr><w:footnoteRef/></w:r>',
        f'<w:r><w:rPr>{base}</w:rPr><w:t xml:space="preserve"> </w:t></w:r>',
    ]
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("`") and token.endswith("`"):
            props, body = f'<w:rFonts w:ascii="{MONO_FONT}" w:hAnsi="{MONO_FONT}"/>{base}', token[1:-1]
        elif token.startswith("**") and token.endswith("**"):
            props, body = f"<w:b/>{base}", token[2:-2]
        elif token.startswith("*") and token.endswith("*"):
            props, body = f"<w:i/>{base}", token[1:-1]
        elif token.startswith("[^"):
            # Una nota dentro una nota non e' rappresentabile: si lascia il
            # testo cosi' com'e' invece di perderlo.
            props, body = base, token
        else:
            props, body = base, token
        runs.append(
            f"<w:r><w:rPr>{props}</w:rPr>"
            f'<w:t xml:space="preserve">{xml_escape(body)}</w:t></w:r>'
        )
    return (
        f'<w:footnote w:id="{fid}">'
        f'<w:p><w:pPr><w:spacing w:after="40"/><w:jc w:val="both"/></w:pPr>'
        + "".join(runs)
        + "</w:p></w:footnote>"
    )


def append_footnote_reference(paragraph, fid: int) -> None:
    """Aggiunge in coda al paragrafo il rimando alla nota `fid`."""
    run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr><w:rStyle w:val="{FOOTNOTE_REF_STYLE}"/></w:rPr>'
        f'<w:footnoteReference w:id="{fid}"/>'
        f"</w:r>"
    )
    paragraph._p.append(run)


def render_inline(paragraph, text: str, id_map: dict[str, int]) -> int:
    """
    Scrive il testo nel paragrafo spezzandolo in run formattati.
    Restituisce quante note sono state inserite.
    """
    inserted = 0
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = MONO_FONT
        elif token.startswith("**") and token.endswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("*") and token.endswith("*"):
            paragraph.add_run(token[1:-1]).italic = True
        elif token.startswith("[^"):
            local = token[2:-1]
            if local in id_map:
                append_footnote_reference(paragraph, id_map[local])
                inserted += 1
            else:
                paragraph.add_run(token)
        else:
            paragraph.add_run(token)
    return inserted


def find_anchor(document, anchor_text: str):
    """Il paragrafo di intestazione prima del quale si inserisce."""
    for par in document.paragraphs:
        if par.style.name.startswith("Heading") and anchor_text.lower() in par.text.lower():
            return par
    return None


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__,
                                     formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--draft", required=True, help="File Markdown con le sezioni da inserire.")
    parser.add_argument("--diary", default=None, help=f"Path del .docx (default: {DIARY_NAME} in root).")
    parser.add_argument("--anchor", default=DEFAULT_ANCHOR,
                        help=f"Intestazione prima della quale inserire (default: {DEFAULT_ANCHOR!r}).")
    parser.add_argument("--apply", action="store_true",
                        help="Scrive davvero nel .docx. Senza questo flag riporta soltanto.")
    parser.add_argument("--no-backup", action="store_true",
                        help="Non creare il .bak. Sconsigliato.")
    args = parser.parse_args()

    root = Path(__file__).resolve().parent.parent
    draft_path = Path(args.draft)
    if not draft_path.is_absolute():
        draft_path = (root / draft_path).resolve()
    diary_path = Path(args.diary).resolve() if args.diary else (root / DIARY_NAME)

    for path in (draft_path, diary_path):
        if not path.is_file():
            print(f"ERRORE: file non trovato: {path}", file=sys.stderr)
            sys.exit(1)

    blocks, footnotes = parse_draft(draft_path.read_text(encoding="utf-8"))
    sections = [b for b in blocks if b["kind"] == "heading" and b["level"] == 2]
    used = sorted({n for b in blocks for n in re.findall(r"\[\^([0-9]+)\]", b["text"])},
                  key=int)

    document = Document(str(diary_path))
    anchor = find_anchor(document, args.anchor)
    fn_part = footnotes_part(document)

    print(f"Draft   : {draft_path}")
    print(f"Diario  : {diary_path}")
    print(f"Ancora  : {anchor.text!r} ({anchor.style.name})" if anchor
          else f"Ancora  : NON TROVATA per {args.anchor!r}")
    print()
    print(f"Sezioni da inserire : {len(sections)}")
    for block in sections:
        print(f"    {block['text']}")
    print(f"Paragrafi           : {sum(1 for b in blocks if b['kind'] == 'para')}")
    print(f"Sottosezioni H3     : {sum(1 for b in blocks if b['kind'] == 'heading' and b['level'] == 3)}")
    print(f"Note a pie' pagina  : {len(used)} usate, {len(footnotes)} definite")

    orphan_refs = [n for n in used if n not in footnotes]
    orphan_defs = [n for n in footnotes if n not in used]
    if orphan_refs:
        print(f"  ATTENZIONE: rimandi senza definizione: {', '.join(orphan_refs)}")
    if orphan_defs:
        print(f"  ATTENZIONE: definizioni mai citate: {', '.join(orphan_defs)}")

    if anchor is None or fn_part is None:
        print("\nERRORE: manca l'ancora oppure la parte footnotes.xml.", file=sys.stderr)
        sys.exit(1)
    if not sections:
        print("\nNessuna sezione trovata nel draft: niente da inserire.", file=sys.stderr)
        sys.exit(1)

    if not args.apply:
        print("\nNessuna modifica scritta. Rilanciare con --apply per inserire.")
        return

    if not args.no_backup:
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        backup = diary_path.with_suffix(f".{stamp}.bak.docx")
        shutil.copy2(diary_path, backup)
        print(f"\nBackup  : {backup.name}")

    # Note: si allocano prima gli id, cosi' i rimandi nel corpo possono
    # puntarli mentre si costruiscono i paragrafi.
    xml = fn_part._blob.decode("utf-8")
    next_id = next_footnote_id(xml)
    id_map: dict[str, int] = {}
    new_notes: list[str] = []
    for local in used:
        if local not in footnotes:
            continue
        id_map[local] = next_id
        new_notes.append(build_footnote_xml(next_id, footnotes[local]))
        next_id += 1
    if new_notes:
        fn_part._blob = xml.replace("</w:footnotes>", "".join(new_notes) + "</w:footnotes>").encode("utf-8")

    before_paragraphs = len(document.paragraphs)
    before_tables = len(document.tables)

    refs_written = 0
    for block in blocks:
        if block["kind"] == "heading":
            style = f"Heading {min(block['level'], 6)}"
            paragraph = document.add_paragraph(style=style)
        else:
            paragraph = document.add_paragraph(style="Normal")
        refs_written += render_inline(paragraph, block["text"], id_map)
        # Il paragrafo nasce in coda al documento: lo si sposta davanti
        # all'ancora mantenendo l'ordine di lettura del draft.
        anchor._p.addprevious(paragraph._p)

    document.save(str(diary_path))

    # Verifica di ritorno: si rilegge il file appena scritto.
    check = Document(str(diary_path))
    check_fn = footnotes_part(check)
    added = len(check.paragraphs) - before_paragraphs
    print(f"\nInseriti {added} paragrafi e {refs_written} rimandi a nota.")
    print(f"Tabelle: {before_tables} prima, {len(check.tables)} dopo.")
    if len(check.tables) != before_tables:
        print("ATTENZIONE: il numero di tabelle e' cambiato, verificare il documento.",
              file=sys.stderr)
    if check_fn is not None:
        total = len(re.findall(r'<w:footnote[^>]*w:id="\d+"', check_fn._blob.decode("utf-8")))
        print(f"Note nel documento dopo l'inserimento: {total}.")
    print("\nPasso successivo: .\\scripts\\finalize_diary.ps1 per rigenerare il .md e "
          "rivedere il diff.")


if __name__ == "__main__":
    main()
