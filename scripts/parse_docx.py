#!/usr/bin/env python3
"""
parse_docx.py — Parsing token-efficient di file .docx aziendali italiani.

Modalita disponibili (subcomandi):
  skeleton          Solo struttura (heading + conteggi). ~50-200 token per documento.
  sections-preview  Sezioni con estratti (primi 200ch + ultimi 100ch).
  full-section      Contenuto completo di una sezione specifica.
  token-estimate    Stima token totali della cartella per pianificazione.

Caratteristiche:
- Parallelizzazione automatica via multiprocessing
- Modalita incrementale: salta i .docx il cui hash non e cambiato
- Gestione corretta di path Windows con spazi, accenti e percorsi OneDrive
- Riconoscimento heading sia da stile (Heading/Titolo) sia da pattern numerici

Pensato per essere chiamato dal subagente Claude, non interattivamente.
"""

import argparse
import hashlib
import json
import re
import sys
import time
from concurrent.futures import ProcessPoolExecutor, as_completed
from dataclasses import dataclass, field, asdict
from datetime import datetime
from multiprocessing import cpu_count
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.document import Document as _Document
    from docx.oxml.ns import qn
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("ERRORE: manca python-docx. Installa con: pip install python-docx", file=sys.stderr)
    sys.exit(2)


# ============================================================
# Util: riconoscimento heading e stima token
# ============================================================

HEADING_STYLE_PATTERNS = [
    re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE),
    re.compile(r"^Titolo\s*(\d+)$", re.IGNORECASE),
    re.compile(r"^H(\d+)$", re.IGNORECASE),
]

NUMERIC_HEADING_PATTERN = re.compile(r"^\s*(\d+(?:\.\d+){0,4})\.?\s+(.+)")


def heading_level(style_name: str) -> Optional[int]:
    """Restituisce il livello 1-6 se lo stile e un heading, altrimenti None."""
    if not style_name:
        return None
    for pat in HEADING_STYLE_PATTERNS:
        m = pat.match(style_name.strip())
        if m:
            try:
                return min(max(int(m.group(1)), 1), 6)
            except ValueError:
                return None
    return None


def numeric_heading_level(text: str) -> Optional[tuple[int, str]]:
    """Riconosce 'X.Y.Z titolo' come heading anche senza stile."""
    m = NUMERIC_HEADING_PATTERN.match(text.strip())
    if not m:
        return None
    numbering = m.group(1)
    title = m.group(2)
    level = numbering.count(".") + 1
    if level > 6:
        return None
    return level, title


def estimate_tokens(text: str) -> int:
    """1 token ≈ 4 caratteri di testo italiano (stima conservativa)."""
    return max(1, len(text) // 4)


def file_hash(path: Path) -> str:
    """SHA256 dei primi 256 KB del file (sufficiente per identificare modifiche)."""
    h = hashlib.sha256()
    try:
        with path.open("rb") as f:
            h.update(f.read(262144))
        return h.hexdigest()[:16]
    except Exception:
        return "unreadable"


# ============================================================
# Iteratori sul corpo .docx
# ============================================================

def iter_block_items(parent):
    """Itera in ordine paragrafi e tabelle di un Document."""
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        return

    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def extract_table_as_dicts(table: Table) -> list[dict]:
    """Converte una tabella in lista di dict, usando la prima riga come header."""
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    if not rows:
        return []
    header = rows[0]
    data = []
    for r in rows[1:]:
        d = {}
        for i, col in enumerate(header):
            key = col if col else f"col_{i}"
            d[key] = r[i] if i < len(r) else ""
        data.append(d)
    return data


# ============================================================
# Strutture dati
# ============================================================

@dataclass
class Section:
    level: int
    title: str
    paragraphs: list[str] = field(default_factory=list)
    tables: list[list[dict]] = field(default_factory=list)
    word_count: int = 0
    char_count: int = 0
    preview_start: str = ""
    preview_end: str = ""

    @property
    def full_text(self) -> str:
        return "\n\n".join(self.paragraphs)

    def finalize_previews(self):
        text = self.full_text
        self.char_count = len(text)
        self.word_count = len(text.split())
        self.preview_start = text[:200].strip()
        self.preview_end = text[-100:].strip() if len(text) > 300 else ""


@dataclass
class DocSkeleton:
    file_path: str
    file_name: str
    relative_path: str
    file_hash: str
    file_size_bytes: int
    mtime_iso: str
    total_paragraphs: int
    total_words: int
    total_chars: int
    estimated_tokens: int
    sections: list[dict]
    tables_count: int
    has_images: bool
    error: Optional[str] = None


# ============================================================
# Parsing di un singolo .docx
# ============================================================

def parse_single_docx(path: Path, base_dir: Path, mode: str) -> DocSkeleton:
    """Parsa un .docx in modalita skeleton / sections-preview / full."""
    try:
        rel = str(path.relative_to(base_dir))
    except ValueError:
        rel = path.name

    try:
        doc = Document(str(path))
    except Exception as e:
        return DocSkeleton(
            file_path=str(path), file_name=path.name, relative_path=rel,
            file_hash="error",
            file_size_bytes=path.stat().st_size if path.exists() else 0,
            mtime_iso=datetime.fromtimestamp(path.stat().st_mtime).isoformat() if path.exists() else "",
            total_paragraphs=0, total_words=0, total_chars=0, estimated_tokens=0,
            sections=[], tables_count=0, has_images=False,
            error=f"{type(e).__name__}: {e}",
        )

    sections: list[Section] = []
    current = Section(level=0, title="Documento completo")
    total_paragraphs = 0
    tables_count = 0
    has_images = False

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text or ""
            if not text.strip():
                continue
            total_paragraphs += 1

            lvl = heading_level(block.style.name if block.style else "")
            if lvl is None:
                numeric = numeric_heading_level(text)
                if numeric:
                    lvl, text = numeric

            if lvl is not None:
                if current.paragraphs or current.tables:
                    sections.append(current)
                current = Section(level=lvl, title=text.strip())
            else:
                current.paragraphs.append(text)

            if "<w:drawing" in block._p.xml or "graphicData" in block._p.xml:
                has_images = True

        elif isinstance(block, Table):
            tables_count += 1
            try:
                current.tables.append(extract_table_as_dicts(block))
            except Exception:
                current.tables.append([])

    if current.paragraphs or current.tables or not sections:
        sections.append(current)

    for s in sections:
        s.finalize_previews()

    total_words = sum(s.word_count for s in sections)
    total_chars = sum(s.char_count for s in sections)

    if mode == "skeleton":
        sec_dicts = [{"level": s.level, "title": s.title,
                      "word_count": s.word_count, "char_count": s.char_count} for s in sections]
    elif mode == "sections-preview":
        sec_dicts = [{"level": s.level, "title": s.title,
                      "word_count": s.word_count, "char_count": s.char_count,
                      "preview_start": s.preview_start, "preview_end": s.preview_end,
                      "tables_count": len(s.tables)} for s in sections]
    elif mode == "full":
        sec_dicts = [{"level": s.level, "title": s.title,
                      "word_count": s.word_count, "char_count": s.char_count,
                      "paragraphs": s.paragraphs, "tables": s.tables} for s in sections]
    else:
        raise ValueError(f"Modalita sconosciuta: {mode}")

    return DocSkeleton(
        file_path=str(path), file_name=path.name, relative_path=rel,
        file_hash=file_hash(path),
        file_size_bytes=path.stat().st_size,
        mtime_iso=datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
        total_paragraphs=total_paragraphs,
        total_words=total_words, total_chars=total_chars,
        estimated_tokens=estimate_tokens("\n".join(p for s in sections for p in s.paragraphs)),
        sections=sec_dicts, tables_count=tables_count, has_images=has_images,
    )


# ============================================================
# Discovery e parallelizzazione
# ============================================================

def find_docx_files(input_path: Path) -> list[Path]:
    """Trova tutti i .docx in una cartella ricorsivamente. Esclude temporanei (~$)."""
    if input_path.is_file():
        return [input_path] if input_path.suffix.lower() == ".docx" else []
    return sorted(
        p for p in input_path.rglob("*.docx")
        if p.is_file() and not p.name.startswith("~$")
    )


def _worker(args_tuple):
    """Wrapper per ProcessPoolExecutor."""
    path_str, base_str, mode = args_tuple
    return asdict(parse_single_docx(Path(path_str), Path(base_str), mode))


def parallel_parse(files: list[Path], base_dir: Path, mode: str,
                   workers: int = 0, log_each: bool = True) -> list[dict]:
    """Parsa in parallelo. workers=0 -> auto (min(cpu, 8))."""
    if workers <= 0:
        workers = min(cpu_count(), 8)
    workers = max(1, min(workers, len(files)))

    if workers == 1:
        results = []
        for i, f in enumerate(files, 1):
            if log_each:
                print(f"  [{i}/{len(files)}] {f.name}", file=sys.stderr)
            results.append(asdict(parse_single_docx(f, base_dir, mode)))
        return results

    print(f"  Parallelismo: {workers} worker", file=sys.stderr)
    args = [(str(f), str(base_dir), mode) for f in files]
    results = []
    start = time.time()
    with ProcessPoolExecutor(max_workers=workers) as ex:
        futures = {ex.submit(_worker, a): i for i, a in enumerate(args)}
        completed = 0
        for fut in as_completed(futures):
            completed += 1
            try:
                results.append(fut.result())
                if log_each and (completed % 10 == 0 or completed == len(files)):
                    elapsed = time.time() - start
                    rate = completed / elapsed if elapsed > 0 else 0
                    eta = (len(files) - completed) / rate if rate > 0 else 0
                    print(f"  [{completed}/{len(files)}] elaborati ({rate:.1f}/s, ETA {eta:.0f}s)",
                          file=sys.stderr)
            except Exception as e:
                print(f"  ! errore: {e}", file=sys.stderr)
    return results


# ============================================================
# Incrementale: carica precedente e filtra invariati
# ============================================================

def load_previous_hashes(output_path: Path) -> dict[str, str]:
    """Legge gli hash dei documenti gia processati da un structure.json esistente."""
    if not output_path.exists():
        return {}
    try:
        prev = json.loads(output_path.read_text(encoding="utf-8"))
        return {doc["file_name"]: doc["file_hash"] for doc in prev.get("documents", [])}
    except Exception:
        return {}


def filter_changed_files(files: list[Path], prev_hashes: dict[str, str]) -> tuple[list[Path], int]:
    """Restituisce solo i file con hash diverso da prima."""
    changed = []
    unchanged = 0
    for f in files:
        new_hash = file_hash(f)
        if prev_hashes.get(f.name) != new_hash:
            changed.append(f)
        else:
            unchanged += 1
    return changed, unchanged


# ============================================================
# Comandi CLI
# ============================================================

def cmd_skeleton(args):
    input_path = Path(args.input).resolve()
    output_path = Path(args.output).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if not input_path.exists():
        print(f"Cartella sorgente non trovata: {input_path}", file=sys.stderr)
        sys.exit(1)

    files = find_docx_files(input_path)
    if not files:
        print(f"Nessun .docx trovato in {input_path}", file=sys.stderr)
        sys.exit(1)

    print(f"Trovati {len(files)} file .docx in {input_path}", file=sys.stderr)
    print(f"Modalita: skeleton", file=sys.stderr)

    prev_documents = []
    if args.incremental:
        prev_hashes = load_previous_hashes(output_path)
        if prev_hashes:
            try:
                prev_documents = json.loads(output_path.read_text(encoding="utf-8")).get("documents", [])
            except Exception:
                prev_documents = []
            files_to_parse, unchanged = filter_changed_files(files, prev_hashes)
            print(f"Incrementale: {unchanged} invariati, {len(files_to_parse)} da rielaborare",
                  file=sys.stderr)
            files = files_to_parse
        else:
            print(f"Incrementale: nessun precedente trovato, eseguo full parsing", file=sys.stderr)

    if not files:
        print("Nessun file da elaborare (tutti invariati).", file=sys.stderr)
        # Riscrive lo stesso file con timestamp aggiornato
        output_path.write_text(
            json.dumps({
                "generated_at": datetime.now().isoformat(),
                "source_root": str(input_path),
                "document_count": len(prev_documents),
                "total_words": sum(d["total_words"] for d in prev_documents),
                "total_estimated_tokens": sum(d["estimated_tokens"] for d in prev_documents),
                "documents": prev_documents,
            }, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        return

    new_results = parallel_parse(files, input_path, mode="skeleton", workers=args.workers)

    # Se incrementale: unisci con i precedenti non rielaborati
    if args.incremental and prev_documents:
        new_names = {r["file_name"] for r in new_results}
        kept = [d for d in prev_documents if d["file_name"] not in new_names]
        all_results = kept + new_results
    else:
        all_results = new_results

    # Rimuovi documenti che non esistono piu sul disco
    current_names = {f.name for f in find_docx_files(input_path)}
    all_results = [d for d in all_results if d["file_name"] in current_names]

    all_results.sort(key=lambda d: d["relative_path"])

    output = {
        "generated_at": datetime.now().isoformat(),
        "source_root": str(input_path),
        "document_count": len(all_results),
        "total_words": sum(r["total_words"] for r in all_results),
        "total_estimated_tokens": sum(r["estimated_tokens"] for r in all_results),
        "documents": all_results,
    }
    output_path.write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\nScritto: {output_path}", file=sys.stderr)
    print(f"Documenti totali: {output['document_count']}", file=sys.stderr)
    print(f"Token stimati totali: {output['total_estimated_tokens']:,}", file=sys.stderr)


def cmd_sections_preview(args):
    input_path = Path(args.input).resolve()
    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    files = find_docx_files(input_path)
    print(f"Trovati {len(files)} file. Modalita: sections-preview", file=sys.stderr)

    if args.incremental:
        # Salta i file il cui JSON corrispondente esiste e ha hash uguale
        files_to_parse = []
        skipped = 0
        for f in files:
            safe_name = re.sub(r"[^\w\-_.]", "_", f.stem) + ".json"
            json_path = output_dir / safe_name
            if json_path.exists():
                try:
                    existing = json.loads(json_path.read_text(encoding="utf-8"))
                    if existing.get("file_hash") == file_hash(f):
                        skipped += 1
                        continue
                except Exception:
                    pass
            files_to_parse.append(f)
        print(f"Incrementale: {skipped} invariati saltati, {len(files_to_parse)} da rielaborare",
              file=sys.stderr)
        files = files_to_parse

    if not files:
        print("Nessun file da elaborare.", file=sys.stderr)
        return

    results = parallel_parse(files, input_path, mode="sections-preview", workers=args.workers)

    for result in results:
        safe_name = re.sub(r"[^\w\-_.]", "_", Path(result["file_name"]).stem) + ".json"
        (output_dir / safe_name).write_text(
            json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8"
        )

    print(f"\nScritti {len(results)} JSON in {output_dir}", file=sys.stderr)


def cmd_full_section(args):
    file_path = Path(args.file).resolve()
    if not file_path.exists():
        print(f"File non trovato: {file_path}", file=sys.stderr)
        sys.exit(1)

    result = parse_single_docx(file_path, file_path.parent, mode="full")
    target = args.section.strip().lower()

    found = None
    for s in result.sections:
        if s["title"].strip().lower() == target or target in s["title"].strip().lower():
            found = s
            break

    if not found:
        print(f"Sezione '{args.section}' non trovata. Disponibili:", file=sys.stderr)
        for s in result.sections:
            print(f"  - {s['title']}", file=sys.stderr)
        sys.exit(1)

    print(json.dumps(found, ensure_ascii=False, indent=2))


def cmd_token_estimate(args):
    input_path = Path(args.input).resolve()
    files = find_docx_files(input_path)
    print(f"Stima token su {len(files)} file...", file=sys.stderr)

    rows = []
    total = 0
    for f in files:
        try:
            doc = Document(str(f))
            text = "\n".join(p.text for p in doc.paragraphs)
            tokens = estimate_tokens(text)
            total += tokens
            rows.append((f.name, tokens, len(text)))
        except Exception as e:
            print(f"  ! errore su {f.name}: {e}", file=sys.stderr)

    rows.sort(key=lambda r: r[1], reverse=True)
    print(f"\n{'File':<60} {'Token':>10} {'Chars':>10}")
    print("-" * 84)
    for name, tok, ch in rows[:30]:
        print(f"{name[:58]:<60} {tok:>10,} {ch:>10,}")
    if len(rows) > 30:
        print(f"... e altri {len(rows) - 30} file")
    print("-" * 84)
    print(f"{'TOTALE':<60} {total:>10,}")
    print(f"\nIn una finestra Sonnet da 200k token, caricarli tutti = {(total/200000)*100:.1f}% del budget.")
    print("Workflow consigliato (scheletro + preview): ~3-5% del budget.")


# ============================================================
# Main
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="Parsing token-efficient di .docx italiani.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    sub = parser.add_subparsers(dest="cmd", required=True, metavar="<comando>")

    common_workers = lambda p: p.add_argument(
        "--workers", type=int, default=0,
        help="Numero di processi paralleli (0 = auto, max 8)"
    )

    p1 = sub.add_parser("skeleton", help="Solo struttura: heading + conteggi")
    p1.add_argument("--input", required=True, help="Cartella con i .docx (ricorsivo)")
    p1.add_argument("--output", required=True, help="Path output JSON")
    p1.add_argument("--incremental", action="store_true",
                    help="Salta i file con hash invariato rispetto a una precedente esecuzione")
    common_workers(p1)
    p1.set_defaults(func=cmd_skeleton)

    p2 = sub.add_parser("sections-preview", help="Sezioni con estratti")
    p2.add_argument("--input", required=True)
    p2.add_argument("--output-dir", required=True, help="Cartella output (un JSON per docx)")
    p2.add_argument("--incremental", action="store_true",
                    help="Salta i .docx il cui JSON esistente ha hash uguale")
    common_workers(p2)
    p2.set_defaults(func=cmd_sections_preview)

    p3 = sub.add_parser("full-section", help="Contenuto completo di una sezione")
    p3.add_argument("--file", required=True)
    p3.add_argument("--section", required=True)
    p3.set_defaults(func=cmd_full_section)

    p4 = sub.add_parser("token-estimate", help="Stima token totale della cartella")
    p4.add_argument("--input", required=True)
    p4.set_defaults(func=cmd_token_estimate)

    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
