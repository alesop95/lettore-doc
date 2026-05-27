"""
sync_diary_md.py - Rigenera il diario .md a partire dal .docx aggiornato.

Mantiene il file Markdown sincronizzato con il documento Word, permettendo
a Git di mostrare diff testuali leggibili delle modifiche.

USO:
    .\.venv\Scripts\python.exe scripts\sync_diary_md.py

INPUT atteso (nella root del repository):
    diario-tecnico-progetto (lettore-doc + skills-repo).docx

OUTPUT (sovrascritto a ogni esecuzione):
    diario-tecnico-progetto (lettore-doc + skills-repo).md
    diario-assets/    (immagini estratte dal .docx)

WORKFLOW:
    1. Modificare il .docx in Word
    2. Eseguire questo script per rigenerare il .md
    3. git add del .docx + .md + diario-assets/
    4. git commit + git push

Il .md include le footnote del .docx come riferimenti [^N] con definizioni
in coda, blocchi di codice fenced, tabelle Markdown standard.
"""
import os
import re
import sys
import zipfile
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# ── Path di default (relativi alla root del repository) ───────────────────
SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parent
DOCX = REPO_ROOT / "diario-tecnico-progetto (lettore-doc + skills-repo).docx"
OUT_MD = REPO_ROOT / "diario-tecnico-progetto (lettore-doc + skills-repo).md"
ASSETS_DIR = REPO_ROOT / "diario-assets"


def extract_footnotes_and_images(docx_path, assets_dir):
    """Estrae note a piè di pagina e immagini dal .docx."""
    footnotes = {}
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    with zipfile.ZipFile(docx_path) as z:
        if "word/footnotes.xml" in z.namelist():
            root = etree.fromstring(z.read("word/footnotes.xml"))
            for fn in root.findall('w:footnote', ns):
                fid = fn.get(qn('w:id'))
                ftype = fn.get(qn('w:type'))
                if ftype in ('separator', 'continuationSeparator'):
                    continue
                texts = []
                for p in fn.findall('.//w:p', ns):
                    run_texts = [t.text or '' for t in p.findall('.//w:t', ns)]
                    texts.append(''.join(run_texts).strip())
                text = ' '.join(t for t in texts if t).strip()
                if text:
                    footnotes[fid] = text

        # Estrai immagini in assets_dir
        assets_dir.mkdir(parents=True, exist_ok=True)
        for f in z.namelist():
            if f.startswith("word/media/"):
                name = os.path.basename(f)
                (assets_dir / name).write_bytes(z.read(f))

    return footnotes


def paragraph_to_md(p, footnotes, fn_used, fn_order):
    """Converte un paragrafo .docx in Markdown, gestendo footnote inline."""
    style = p.style.name if p.style else "Normal"
    parts = []
    for elem in p._p.iter():
        tag = etree.QName(elem.tag).localname
        if tag == 't' and elem.text:
            parts.append(elem.text)
        elif tag == 'footnoteReference':
            fid = elem.get(qn('w:id'))
            if fid in footnotes:
                if fid not in fn_used:
                    fn_used[fid] = len(fn_used) + 1
                    fn_order.append(fid)
                parts.append(f"[^{fn_used[fid]}]")
        elif tag == 'tab':
            parts.append('\t')
        elif tag == 'br':
            parts.append('\n')

    text = ''.join(parts).strip()
    if not text:
        return None

    if style == "Heading 1":
        return f"\n# {text}\n"
    if style == "Heading 2":
        return f"\n## {text}\n"
    if style == "Heading 3":
        return f"\n### {text}\n"
    if style == "Heading 4":
        if re.match(r'^[.…]+$', text):
            return None  # salta H4 vuoti tipo "..."
        return f"\n#### {text}\n"
    if style == "Heading 5":
        return f"\n##### {text}\n"
    return f"{text}\n"


def is_code_block(table):
    """Una tabella è blocco di codice se è 1x1 e usa font monospace."""
    if len(table.rows) != 1 or len(table.columns) != 1:
        return False
    cell = table.cell(0, 0)
    for p in cell.paragraphs:
        for run in p.runs:
            fname = run.font.name or ''
            if 'Consolas' in fname or 'Courier' in fname or 'Mono' in fname:
                return True
    return False


def code_block_to_md(table):
    """Estrae le righe di una tabella code-block come fenced markdown."""
    cell = table.cell(0, 0)
    lines = [p.text for p in cell.paragraphs]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    return "\n```\n" + "\n".join(lines) + "\n```\n"


def table_to_md(table):
    """Converte una tabella dati in tabella Markdown."""
    rows = []
    for r in table.rows:
        cells = [
            c.text.strip().replace('\n', ' ').replace('|', '\\|')
            for c in r.cells
        ]
        rows.append(cells)
    if not rows:
        return ""
    n = len(rows[0])
    md = "\n| " + " | ".join(rows[0]) + " |\n"
    md += "|" + "|".join(["---"] * n) + "|\n"
    for r in rows[1:]:
        while len(r) < n:
            r.append("")
        md += "| " + " | ".join(r) + " |\n"
    return md


def convert(docx_path, out_md_path, assets_dir):
    """Conversione completa .docx -> .md."""
    if not docx_path.exists():
        print(f"ERRORE: .docx non trovato: {docx_path}", file=sys.stderr)
        sys.exit(1)

    footnotes = extract_footnotes_and_images(docx_path, assets_dir)
    doc = Document(str(docx_path))

    fn_used = {}
    fn_order = []
    output = []
    in_body = False

    for child in doc.element.body.iterchildren():
        tag = etree.QName(child.tag).localname
        if tag == 'p':
            p_obj = next((p for p in doc.paragraphs if p._p is child), None)
            if p_obj is None:
                continue
            # Salta la copertina prima di "# Introduzione"
            if not in_body:
                if p_obj.style.name == "Heading 1" and "Introduzione" in (p_obj.text or ""):
                    in_body = True
                else:
                    continue
            md = paragraph_to_md(p_obj, footnotes, fn_used, fn_order)
            if md:
                output.append(md)
        elif tag == 'tbl' and in_body:
            t_obj = next((t for t in doc.tables if t._tbl is child), None)
            if t_obj is None:
                continue
            if is_code_block(t_obj):
                output.append(code_block_to_md(t_obj))
            else:
                output.append(table_to_md(t_obj))

    if fn_order:
        output.append("\n\n---\n\n## Note\n\n")
        for fid in fn_order:
            num = fn_used[fid]
            text = re.sub(r'\s+', ' ', footnotes[fid]).strip()
            output.append(f"[^{num}]: {text}\n\n")

    md = "".join(output)
    # Pulizia spaziatura
    md = re.sub(r'\n{3,}', '\n\n', md)
    md = re.sub(
        r'([a-zàèéìòùA-Z0-9\)\.\!\?\"\'])\n([A-ZÀ-Ý#\|`\[])',
        r'\1\n\n\2', md
    )
    md = re.sub(r'\n{3,}', '\n\n', md)

    header = (
        "# Tassonomia di Competenze IT — Diario tecnico di sviluppo\n\n"
        "*Versione Markdown del documento "
        "`diario-tecnico-progetto (lettore-doc + skills-repo).docx`. "
        "Le due versioni sono mantenute sincronizzate: vedere `CLAUDE.md` "
        "di lettore-doc per le regole di aggiornamento.*\n\n"
        "**Repository**: `alesop95/lettore-doc` (privato) + "
        "`alesop95/skills` (pubblico)  \n"
        "**Sito live**: [alesop95.github.io/skills/]"
        "(https://alesop95.github.io/skills/)\n\n"
        "---\n"
    )

    out_md_path.write_text(header + md.lstrip(), encoding='utf-8')

    print(f"Convertito: {docx_path.name}")
    print(f"      -> {out_md_path.name}")
    print(f"Dimensione output: {out_md_path.stat().st_size:,} bytes")
    print(f"Note a pie' di pagina convertite: {len(fn_used)}")
    print(f"Immagini estratte in: {assets_dir.name}/")


if __name__ == "__main__":
    convert(DOCX, OUT_MD, ASSETS_DIR)
